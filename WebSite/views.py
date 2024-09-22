from django.shortcuts import render,get_object_or_404
from django.http import HttpResponse
from .models import *
from PyPDF2 import *
import io
from docx import Document

from docx import Document
# Create your views here.
def home(request):
    # Uncomment this if you plan to render a template
    # return render(request, 'home.html', {})
    #return HttpResponse("Hi there, this is a Django response")
    subjects = SubjectModel.objects.all()
    return render(request,'index.html',{'subjects': subjects})

def demo(request):
    # Uncomment this if you plan to render a template
    # return render(request, 'home.html', {})
    #return HttpResponse("Hi there, this is a Django response"

    subjects = SubjectModel.objects.all()
    #clicked_button_name=request.GET.get('clicked_button')
    return render(request,'index.html',{'subjects':subjects})

def admin01(request):
    return render(request,'admin.html',{})

def mypdf(request):
    return render(request,'pdf.html',{})

def deletpage(request):
    if request.method == 'POST':
        pdf_file = request.FILES.get('uploadPDF')
        pages_to_keep = request.POST.get('pages')

        if not pdf_file or not pages_to_keep:
            return HttpResponse('Please upload a PDF and specify pages to keep.', status=400)

        try:
            # Convert comma-separated page numbers into a list of integers (user-friendly, so 1-based index)
            pages_to_keep = [int(page.strip()) - 1 for page in pages_to_keep.split(',')]
        except ValueError:
            return HttpResponse("Invalid page numbers format.", status=400)

        # Process the PDF using PyPDF2
        reader = PdfReader(pdf_file)
        writer = PdfWriter()

        total_pages = len(reader.pages)

        # Add only the pages the user wants to keep
        for page_num in pages_to_keep:
            if 0 <= page_num < total_pages:
                writer.add_page(reader.pages[page_num])
            else:
                return HttpResponse(f"Page number {page_num + 1} is out of range.", status=400)

        # Create an in-memory buffer to save the modified PDF
        output_pdf = io.BytesIO()
        writer.write(output_pdf)
        output_pdf.seek(0)  # Move to the beginning of the BytesIO object

        # Send the edited PDF as a response
        response = HttpResponse(output_pdf, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="edited_pdf.pdf"'
        return response

    # For GET request, just render the HTML form
    return render(request, 'Cut_Pdf.html')


import io
from PyPDF2 import PdfReader, PdfWriter
from django.http import HttpResponse

def merge_pdfs(request):
    if request.method == 'POST':
        pdf1_file = request.FILES.get('pdf1')
        pdf2_file = request.FILES.get('pdf2')

        if not pdf1_file or not pdf2_file:
            return HttpResponse('Please upload both PDF files.', status=400)

        writer = PdfWriter()

        # Read the first PDF file
        pdf1_reader = PdfReader(pdf1_file)
        for page in pdf1_reader.pages:
            writer.add_page(page)

        # Read the second PDF file
        pdf2_reader = PdfReader(pdf2_file)
        for page in pdf2_reader.pages:
            writer.add_page(page)

        # Create an in-memory buffer to save the merged PDF
        output_pdf = io.BytesIO()
        writer.write(output_pdf)
        output_pdf.seek(0)  # Move to the beginning of the BytesIO object

        # Send the merged PDF as a response
        response = HttpResponse(output_pdf, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="merged_pdf.pdf"'
        return response

    # For GET request, just render the HTML form (you can customize as needed)
    return render(request, 'merge_pdf.html')



def pdf_to_word(request):
    if request.method == 'POST':
        pdf_file = request.FILES.get('uploadPDF')

        if not pdf_file:
            return HttpResponse('Please upload a PDF file.', status=400)

        # Create a new Document
        doc = Document()

        # Read the PDF file
        reader = PdfReader(pdf_file)

        # Extract text from each page and add to the Word document
        for page in reader.pages:
            text = page.extract_text()
            if text:
                doc.add_paragraph(text)
            doc.add_paragraph()  # Add a new paragraph for spacing between pages

        # Create an in-memory buffer to save the Word document
        output_word = io.BytesIO()
        doc.save(output_word)
        output_word.seek(0)  # Move to the beginning of the BytesIO object

        # Send the Word document as a response
        response = HttpResponse(output_word, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="converted_word.docx"'
        return response

    # For GET request, just render the HTML form (customize as needed)
    return render(request, 'pdf_to_word.html')



import os
from django.http import HttpResponse
from django.core.files.storage import FileSystemStorage
from docx2pdf import convert
import os
import tempfile
from django.http import HttpResponse
from django.core.files.storage import FileSystemStorage
from docx2pdf import convert

def word_to_pdf(request):
    if request.method == 'POST':
        word_file = request.FILES.get('uploadWord')

        if not word_file:
            return HttpResponse('Please upload a Word file.', status=400)

        # Create a temporary directory
        with tempfile.TemporaryDirectory() as temp_dir:
            # Save the uploaded file to the temporary directory
            fs = FileSystemStorage(location=temp_dir)
            filename = fs.save(word_file.name, word_file)
            file_path = os.path.join(temp_dir, filename)

            # Convert the Word file to PDF
            pdf_path = os.path.join(temp_dir, f"{os.path.splitext(filename)[0]}.pdf")
            convert(file_path, pdf_path)

            # Prepare the PDF file for download
            with open(pdf_path, 'rb') as pdf_file:
                response = HttpResponse(pdf_file.read(), content_type='application/pdf')
                response['Content-Disposition'] = f'attachment; filename="{os.path.basename(pdf_path)}"'

            return response

    # For GET request, just render the HTML form (customize as needed)
    return render(request, 'word_to_pdf.html')

